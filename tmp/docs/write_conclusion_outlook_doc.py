from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


REFERENCE_DOC = Path(r"C:\Users\35108\Desktop\总结与展望.docx")
TARGET_DOC = Path(r"C:\Users\35108\Desktop\车险理赔预测系统_总结与展望.docx")


NEW_PARAGRAPHS = [
    "第 6 章 总结与展望  ",
    "6.1 主要工作 ",
    "本文所完成的主要工作可概括为两个方面。第一部分是面向机动车保险业务的理赔率预测模型设计与实现，第二部分是车险理赔预测系统的设计与开发。在模型部分，围绕“保单在当前年度是否发生理赔”这一二分类任务，完成了训练数据构建、日期锚点特征工程、标准化预处理、残差式MLP主干网络设计、带正类权重的BCEWithLogitsLoss损失函数、AdamW优化、Cosine Warmup学习率调度、AUC监控下的Early Stopping以及自动阈值搜索等关键环节的实现，并完成了模型训练、验证、测试及在线推理部署。 ",
    "在系统部分，构建了一个基于Spring Boot、Vue、MySQL和FastAPI的前后端分离车险理赔预测平台，实现了用户登录、保单信息管理、车辆信息管理、理赔记录管理、统计分析、预测管理、预测统计以及模型训练等功能模块。系统能够对业务数据进行统一存储、查询与维护，并支持从保单库中选择记录发起风险预测，完成“数据管理—模型预测—结果分析”的完整流程。与此同时，系统还结合可视化页面对保费、理赔记录、风险等级及预测结果进行直观展示，从而为车险业务中的风险识别和辅助决策提供了较为完整的实现方案。 ",
    "综上所述，本文完成了车险理赔预测模型与业务系统的结合，实现了从模型研究到系统落地的基本闭环。 ",
    "6.2 研究的不足  ",
    "本文的不足主要体现在以下几个方面。首先，数据来源和特征类型仍然较为有限。现有模型主要依赖保单基础信息、车辆属性和历史理赔记录等结构化字段，尚未融合驾驶行为、道路环境、天气因素或更长时间跨度的动态业务数据，因此模型对复杂风险模式的刻画能力仍有一定限制。其次，模型任务目前聚焦于“是否发生理赔”的单任务分类预测，虽然能够较好地支持风险等级识别，但在理赔金额估计、理赔类型细分以及多目标联合建模等方面仍缺乏进一步拓展。 ",
    "再次，系统虽然已经实现了数据管理、统计分析、在线预测与训练管理等核心功能，但在业务闭环方面仍有提升空间。例如，系统尚未形成更完善的预警推送、报表导出、细粒度权限控制以及与真实保险业务流程深度联动的机制，模型解释结果与业务决策规则之间也仍可进一步加强。上述问题说明，本文所完成的工作已经具备较好的研究和实现基础，但在数据广度、模型深度和系统完整性方面仍存在后续优化空间。 ",
    "6.3 未来的展望  ",
    "未来的研究可从模型能力与数据质量两个方向进一步展开。在模型层面，可在现有残差式MLP分类框架的基础上继续探索更适合保险表格数据的结构设计，例如引入更细粒度的特征交互机制、开展多任务学习，或者结合可解释性方法增强模型输出对业务人员的可理解程度。在数据层面，可进一步扩充训练样本的时间跨度与业务覆盖范围，并引入外部环境数据、用户行为数据以及更多与风险相关的上下文特征，从而提高模型在复杂场景下的泛化能力与稳定性。 ",
    "在系统应用层面，未来可继续完善车险理赔预测平台的工程化能力，例如增强模型版本管理与自动评估机制，完善批量预测、报表生成、预警提醒和结果追踪等功能，并推动预测结果与续保定价、风险分层、人工审核等业务环节形成更紧密的联动。随着人工智能技术在保险行业中的不断深化应用，车险理赔预测系统有望从单纯的预测工具逐步发展为集数据管理、风险识别、辅助决策和运营支持于一体的综合性智能平台。 ",
]


def set_songti(paragraph, size_pt: float = 12) -> None:
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size_pt)


def set_heading_style(paragraph) -> None:
    set_songti(paragraph, 14)
    for run in paragraph.runs:
        run.bold = True


def main() -> None:
    if not REFERENCE_DOC.exists():
        raise FileNotFoundError(f"Reference DOCX not found: {REFERENCE_DOC}")

    doc = Document(REFERENCE_DOC)

    for i, text in enumerate(NEW_PARAGRAPHS):
        if i < len(doc.paragraphs):
            paragraph = doc.paragraphs[i]
            paragraph.text = text
        else:
            paragraph = doc.add_paragraph(text)

        if i in {0, 1, 5, 8}:
            set_heading_style(paragraph)
        else:
            set_songti(paragraph, 12)

    for j in range(len(NEW_PARAGRAPHS), len(doc.paragraphs)):
        doc.paragraphs[j].text = ""

    doc.save(TARGET_DOC)
    print(f"UPDATED={TARGET_DOC}")
    print(f"REFERENCE={REFERENCE_DOC}")


if __name__ == "__main__":
    main()
