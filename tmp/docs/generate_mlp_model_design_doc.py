from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"D:\Gradua_pro_code-spring\Vehic_Insur_Claim")
OUTPUT_PATH = ROOT / "output" / "doc" / "MLP模型设计说明.docx"


def set_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for style_name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 13)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "模块"
    header[1].text = "默认设计"
    header[2].text = "作用说明"

    rows = [
        ("主干隐藏层", "(256, 512, 512, 256, 256)", "逐层提取共享表征，先扩维再回收维度"),
        ("分类头隐藏层", "64", "在共享表征上做二分类判别"),
        ("输入 Dropout", "0.0", "当前默认不在输入端做随机失活"),
        ("主干 Dropout", "0.25", "控制主干网络过拟合"),
        ("头部 Dropout", "0.15", "控制分类头过拟合"),
        ("头部采样数", "1", "支持多次 dropout 采样后平均 logit"),
        ("正类权重", "4.10", "缓解类别不平衡对正类学习的不利影响"),
        ("标签平滑", "0.0", "默认关闭，可选减弱过度自信"),
    ]

    for module_name, design, desc in rows:
        cells = table.add_row().cells
        cells[0].text = module_name
        cells[1].text = design
        cells[2].text = desc


def build_document() -> Document:
    doc = Document()
    set_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("MLP 模型设计说明")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run("基于 MLP/Model.py，参考 TrainConfig.py 与 TrainModel.py")

    doc.add_paragraph("")

    doc.add_heading("1. 文档范围", level=1)
    add_paragraph(
        doc,
        "本文只解释当前项目中 MLP 模型本身的设计，不展开训练流程、优化器、学习率调度和 FastAPI 接口部分。"
        "分析依据主要来自 MLP/Model.py，同时参考 TrainConfig.py 中的默认超参数，以及 TrainModel.py 中模型实例化和损失参数解析逻辑。"
    )
    add_paragraph(
        doc,
        "从 TrainModel.py 的 build_model_config_payload 可以看出，当前模型被明确标记为 "
        "architecture = claim_classifier_residual_mlp，task_type = classification_only。"
        "这说明系统当前已经从早期的多任务或更复杂兼容结构，收敛到一个以‘是否出险/是否理赔’为目标的单任务二分类模型。"
    )

    doc.add_heading("2. 总体模型框架", level=1)
    add_paragraph(
        doc,
        "当前主模型类是 InsuranceMLP。它的整体计算路径可以概括为：输入表格特征 x -> "
        "输入端 Dropout -> 共享主干 SharedBackbone -> 分类头 ClassificationHead -> 输出单个分类 logit。"
        "在推理阶段，logit 会再通过 sigmoid 映射为理赔概率，并结合阈值产生最终的 0/1 分类结果。"
    )
    add_bullet(doc, "输入层不直接做复杂特征交互，而是先把原始特征投影到第一个隐藏空间。")
    add_bullet(doc, "中间主体采用残差式 MLP 主干，而不是普通的逐层串联全连接网络。")
    add_bullet(doc, "输出端只保留一个二分类头，不再包含回归头等多任务输出。")
    add_bullet(doc, "模型的最终输出是 logit，而不是概率，这与 BCEWithLogitsLoss 的数值稳定实现相配套。")

    doc.add_heading("3. 主干网络 SharedBackbone 的设计", level=1)
    add_paragraph(
        doc,
        "SharedBackbone 由一个输入投影层和若干个残差块组成。默认隐藏维度为 "
        "(256, 512, 512, 256, 256)。因此，网络先把输入映射到 256 维，再经过四个残差块完成表征提取。"
        "这些残差块的维度变化依次为 256->512、512->512、512->256、256->256。"
    )
    add_paragraph(
        doc,
        "输入投影层 input_proj 采用 Linear + LayerNorm + GELU + Dropout 的组合。"
        "这种写法的核心意图是：先完成线性映射，再用 LayerNorm 让每个样本的特征分布更稳定，"
        "最后通过 GELU 提供平滑非线性，并用 dropout 抑制过拟合。与 BatchNorm 相比，LayerNorm 对 batch 大小的依赖更弱，"
        "更适合批量大小可能波动、且特征主要是表格数值输入的场景。"
    )
    add_paragraph(
        doc,
        "每个 ResidualBlock 的内部结构为：fc1 -> norm1 -> GELU -> Dropout -> fc2 -> norm2 -> 与 shortcut 相加 -> GELU。"
        "这里的 shortcut 不是固定恒等映射，而是‘维度相同则 Identity，不同则用一个无偏置线性层投影到目标维度’。"
        "因此，这个残差结构不仅能稳定深层训练，还能允许相邻层之间自由改变通道维度。"
    )
    add_bullet(doc, "特点 1：采用残差连接，降低深层全连接网络训练中的梯度退化风险。")
    add_bullet(doc, "特点 2：残差支路支持维度变换，因此主干不受‘相邻层维度必须一致’的限制。")
    add_bullet(doc, "特点 3：主干全部使用 GELU 和 LayerNorm，风格更接近现代 Transformer/MLP 设计，而不是传统 ReLU + BatchNorm。")
    add_bullet(doc, "特点 4：Dropout 放在第一层激活之后，有助于在高维表示形成早期就引入随机正则。")

    doc.add_heading("4. 分类头 ClassificationHead 的设计", level=1)
    add_paragraph(
        doc,
        "分类头相对轻量，默认结构为 Linear(主干输出维度, 64) -> LayerNorm -> GELU -> Dropout(0.15) -> Linear(64, 1)。"
        "这意味着主干负责学习更通用、更稳定的共享表示，而分类头只完成最终的判别映射。"
        "这种‘重主干、轻头部’的分工对表格二分类是比较自然的：主干负责提炼结构化信息，头部负责做最终决策。"
    )
    add_paragraph(
        doc,
        "需要特别注意的是，InsuranceMLP 提供了 head_samples 参数。若 head_samples > 1，模型会对同一个共享特征重复调用同一个分类头若干次，"
        "并对得到的多个 logit 求平均。由于这些重复调用共享同一套参数，所以真正产生差异的来源主要是 dropout 的随机性。"
        "因此，这个机制本质上更接近一种 Monte Carlo Dropout 风格的头部采样平均，而不是多个独立头的集成。"
    )
    add_paragraph(
        doc,
        "这一设计在训练态下更有意义，因为训练时 dropout 是开启的，多次前向会形成不同的随机子网络；"
        "但在标准推理态下 model.eval() 会关闭 dropout，此时重复调用同一个头部将得到几乎完全相同的输出，"
        "于是 head_samples 的效果会退化成一次确定性前向的重复平均。也就是说，当前默认配置 head_samples = 1 是合理的。"
    )

    doc.add_heading("5. 参数初始化设计", level=1)
    add_paragraph(
        doc,
        "InsuranceMLP 在初始化时会显式遍历模块并重置参数。对所有 Linear 层使用 Kaiming Normal 初始化，偏置置零；"
        "对所有 LayerNorm 的缩放参数置 1、偏置置 0。这样做的目的，是让网络在训练开始时保持较稳定的信号尺度。"
        "虽然代码中 Kaiming 初始化的 nonlinearity 参数写的是 relu，而模型真实激活使用的是 GELU，"
        "两者并不是完全严格匹配，但在实践中通常仍能提供较好的初始化效果。"
    )

    doc.add_heading("6. 损失函数设计", level=1)
    add_paragraph(
        doc,
        "损失函数类是 ClaimClassificationLoss，本质上是对二分类任务的 BCEWithLogitsLoss 做了一层封装。"
        "由于模型输出的是 logit，因此这里直接使用 binary_cross_entropy_with_logits，"
        "避免先 sigmoid 再做 BCE 时可能带来的数值稳定性问题。"
    )
    add_paragraph(
        doc,
        "该损失包含两个关键增强点。第一是 pos_weight：代码把正类权重注册成 buffer，"
        "在计算 BCEWithLogitsLoss 时传入 pos_weight，从而提高正类样本在损失中的权重。"
        "这非常适合处理理赔预测中常见的类别不平衡问题。第二是 label_smoothing："
        "目标标签会按公式 y_smooth = y * (1 - eps) + 0.5 * eps 做平滑。"
        "也就是说，当 eps > 0 时，原本的 0/1 硬标签会向 0.5 稍微收缩，降低模型过度自信的倾向。"
    )
    add_paragraph(
        doc,
        "TrainModel.py 中又进一步说明了 pos_weight 的解析方式：若配置中的 pos_weight <= 0，"
        "则会根据训练集负样本数/正样本数自动估算；否则使用配置值。之后还会强制 pos_weight >= 1.0。"
        "这表明损失函数的设计思路非常明确，即优先服务于不平衡二分类。"
    )
    add_paragraph(
        doc,
        "值得注意的是，ClaimClassificationLoss.forward 最终返回的是 (clf_loss, clf_loss) 这样一个二元组。"
        "从 TrainModel.py 中的调用方式 loss, clf_loss = criterion(logits, labels) 可以看出，"
        "训练器预期同时拿到‘总损失’和‘分类损失’两个值。当前模型只有分类任务，"
        "因此直接把同一个损失同时作为 total loss 和 clf loss 返回。这是一种典型的接口兼容设计，"
        "说明代码曾经为多任务或多损失结构预留过统一训练接口。"
    )

    doc.add_heading("7. 兼容性设计与模型演化痕迹", level=1)
    add_paragraph(
        doc,
        "Model.py 的前半部分定义的是当前主模型，而后半部分保留了若干以 _Compatible 开头的旧结构，"
        "包括 _CompatibleSimpleMLP 和 _CompatibleAdvancedMLP 及其对应的 stem、backbone、head。"
        "这些类并不是当前训练主路径的核心，而是为了兼容旧 checkpoint 的推理加载。"
    )
    add_paragraph(
        doc,
        "build_model_from_checkpoint 会通过 state_dict 中键名的模式来识别模型来源。"
        "如果检测到当前残差主干结构，就优先恢复 InsuranceMLP；"
        "如果检测到旧版 simple 结构或 advanced 结构，就分别构造对应的兼容模型；"
        "如果检测到早期多任务模型留下的 reg_head，也会在加载时把无关的回归头权重过滤掉。"
        "此外，它还能根据 checkpoint 中实际权重张量的形状反推出 hidden_dims 和 head_hidden_dim。"
    )
    add_paragraph(
        doc,
        "这说明模型设计不仅考虑了当前效果，也考虑了系统演进过程中的可维护性。"
        "从工程视角看，这是一种很实用的设计：训练主线已经简化为单任务残差 MLP，"
        "但历史实验和旧模型文件仍然可以被安全加载，不会因为结构升级而完全失效。"
    )

    doc.add_heading("8. 当前模型设计的核心特点总结", level=1)
    add_bullet(doc, "这是一个面向表格数据二分类任务的单任务残差式 MLP，而不是卷积网络或序列模型。")
    add_bullet(doc, "主干强调共享特征提取，分类头保持轻量，符合‘先学表示、后做判别’的设计思想。")
    add_bullet(doc, "通过 LayerNorm + GELU + 残差连接提升深层 MLP 的训练稳定性和表示能力。")
    add_bullet(doc, "通过 pos_weight 与可选 label smoothing 让损失函数更贴合理赔场景中的不平衡分类。")
    add_bullet(doc, "通过 head_samples 预留了基于 dropout 的头部采样平均机制，增强了结构灵活性。")
    add_bullet(doc, "通过兼容加载逻辑保留了历史模型的可恢复性，体现出较强的工程连续性。")

    doc.add_heading("9. 结论", level=1)
    add_paragraph(
        doc,
        "综合来看，当前项目中的模型并不是一个追求极端复杂结构的网络，而是一个围绕表格分类任务精心打磨的残差式 MLP。"
        "它在结构上强调稳定、可训练、可扩展，在损失上强调处理类别不平衡，在工程上强调与历史版本兼容。"
        "如果用一句话概括，这个模型的设计目标就是：用现代化残差 MLP 的方式，为机动车保险理赔风险判别提供一个稳定、清晰、易维护的二分类骨干。"
    )

    doc.add_heading("附：当前默认配置速览", level=1)
    add_table(doc)

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
